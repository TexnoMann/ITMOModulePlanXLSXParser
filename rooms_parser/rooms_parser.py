import pandas as pd
from docx.api import Document
import json, re, datetime

def isInt(number):
    try:
        inNumberfloat = int(number)
    except ValueError:
        return False
    return True

class RoomsParser:
    def __init__(self, input_file, rooms_info, time_config):
        self.__document = Document(input_file)
        self.__tables = self.__document.tables
        self.__tables_dict = {}
        self.__rooms_info = {}
        self.__building_names = {}
        self.__time_period_list = []
        self.__rooms_busy_list = []
        self.__comment_info = "ПК и ФМШ"
        rinfo_file = open(rooms_info)
        tconf_file = open(time_config)
        self.__parseRoomsInfo(rinfo_file)
        self.__parseTimeConfig(tconf_file)


    def __parseRoomsInfo(self, rinfo_file):
        info_list = json.load(rinfo_file)
        for i in range(0, len(info_list)):
            cell = info_list[i]
            if cell['building_id'] not in self.__rooms_info.keys():
                self.__rooms_info[cell['building_id']]={}
            if cell['building_id'] not in self.__building_names.keys():
                self.__building_names[cell['building_id']] = cell['building']
            self.__rooms_info[cell['building_id']][cell['room_number']] = cell['r_id']
        # print(self.__building_names)

    def __parseTimeConfig(self, time_config_file):
        self.__time_config_dict = json.load(time_config_file)
        self.__lesson_time = datetime.datetime.strptime(self.__time_config_dict['lesson_time'],'%H:%M').time()
        for i in range(0,len(self.__time_config_dict['start_times'])):
            self.__time_period_list.append(datetime.datetime.strptime(self.__time_config_dict['start_times'][i],'%H:%M').time())
        # print(self.__time_period_list)


    def parseRoomsTable(self):
        corps = self.getText()
        for i in range(0,len(corps)):
            print("[PROCESSING] : Scan building_name: ", corps[i], ".....")
            keys = None
            table = self.__tables[i]
            data = []
            for j, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                if j == 0:
                    keys = tuple(text)
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)
            df = pd.DataFrame(data)

            header = list(df)
            for r in range(0, len(header)):
                list_split_h = re.split(';|,| ',header[r])
                if len(list_split_h)>1:
                    s_head = list_split_h[0]
                    for j in range(1,len(list_split_h)):
                        df[list_split_h[j]] = df[header[r]]
                    df[s_head] = df[header[r]]
                    df=df.drop(columns=[header[r]])

            self.__tables_dict[corps[i]] = df
            header = list(df)

            for h in header:
                if h!='':
                    building_id = -1
                    for b in self.__building_names.keys():
                        res = re.search(corps[i].split(' ')[1], self.__building_names[b])
                        # print(res)
                        if res != None and  h in self.__rooms_info[b].keys():
                            building_id = b
                            break
                    if(building_id != -1):
                        busy_table_col = df[h]
                        for time_row_index in range(0,len(busy_table_col)):
                            time_row = busy_table_col[time_row_index]
                            # print(len(time_row))
                            if time_row!='':
                                start_finish_times = self.__stringToSFTime(time_row)
                                for k in range(0, len(self.__time_period_list)):
                                    if start_finish_times[0]<=self.__time_period_list[k] and start_finish_times[1]>self.__time_period_list[k] :
                                        rooms_busy_info ={"r_id":self.__rooms_info[building_id][h], "room_number": h, "day": time_row_index, "pairIndex": k+1, "comment": self.__comment_info}
                                        self.__rooms_busy_list.append(rooms_busy_info)
                        print("[PROCESSING] : Write info about room ", h)
                    else:
                        print("[ERROR] : Cannot find building id for:", h)


    def __stringToSFTime(self, input_string):
        split_string = re.split(';|,|-| ',input_string)
        time1 = []
        time2 = []
        counter = 0
        for s in split_string:
            if isInt(s):
                if len(time1)<=1 and len(time2)==0:
                    time1.append(s)
                elif len(time1) == 2 and len(time2)<=1:
                    time2.append(s)
        dt1 = datetime.datetime.strptime(":".join(time1),'%H:%M').time()
        dt2 = datetime.datetime.strptime(":".join(time2),'%H:%M').time()
        return (dt1, dt2)

    def save(self, output_filename):
        json_str_out = json.dumps(self.__rooms_busy_list, indent = 4, ensure_ascii=False)
        f = open(output_filename, "w")
        print("[COMPLETE] : Write json output in: ", output_filename)
        f.write(json_str_out)
        f.close()

    def getText(self):
        fullText = []
        for para in self.__document.paragraphs:
            if para.text != "":
                fullText.append(para.text)
        return fullText
